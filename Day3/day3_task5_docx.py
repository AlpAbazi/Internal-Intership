#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading("Day3 - Task 5: Programimi me Enum dhe Zgjedhje të Kufizuara - Dokumentacioni", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add metadata
metadata = doc.add_paragraph()
metadata.add_run("Data e Përfundimit: ").bold = True
metadata.add_run("16 Prill 2026\n")
metadata.add_run("Studenti: ").bold = True
metadata.add_run("Internal Internship - Week 1\n")
metadata.add_run("Tema: ").bold = True
metadata.add_run("Enum, Zgjedhje të Kufizuara dhe Logjikë Programi me Kushte ose Switch")

doc.add_paragraph()

# 1. PËRSHKRIMI I PROGRAMIT
doc.add_heading("1. PËRSHKRIMI I PROGRAMIT", level=2)

doc.add_heading("Qëllimi", level=3)
doc.add_paragraph("Programi demonstron përdorimin e enum për të përfaqësuar një grup të kufizuar vlerash dhe reagimin ndryshe sipas zgjedhjes së përdoruesit. Modelon nivelin e një nxënësi.")

doc.add_heading("Funksionalitetet Kryesore", level=3)
doc.add_paragraph("1. Enum Level:")
doc.add_paragraph("Beginner, Intermediate, Advanced, Expert", style='List Bullet 2')

doc.add_paragraph("2. Input nga Përdoruesi:")
doc.add_paragraph("Zgjedhje numerike (0-3) për nivelin", style='List Bullet 2')

doc.add_paragraph("3. Logjikë me Switch:")
doc.add_paragraph("Shfaq mesazh të ndryshëm për secilën vlerë enum", style='List Bullet 2')

doc.add_paragraph("4. Trajtimi i Input-it të Pavlefshëm:")
doc.add_paragraph("Mesazh i qartë për zgjedhje jashtë rangut", style='List Bullet 2')

doc.add_paragraph("5. For Loop për Tre Zgjedhje:")
doc.add_paragraph("Kontrollon tre zgjedhje radhazi", style='List Bullet 2')

doc.add_paragraph("6. Numërues dhe Statistika:")
doc.add_paragraph("Numëron zgjedhjet valide dhe të pavlefshme, shfaq statistikat në fund", style='List Bullet 2')

# 2. LOGJIKA E PROGRAMIT
doc.add_heading("2. LOGJIKA E PROGRAMIT", level=2)

doc.add_paragraph("Programi përdor një enum për të definuar katër nivele të nxënësit. Në një cikël for që përsëritet tre herë, merr një zgjedhje numerike nga përdoruesi. Nëse zgjedhja është brenda 0-3, e konverton në enum dhe përdor switch për të shfaqur mesazhin përkatës. Për zgjedhje të pavlefshme, shfaq mesazh gabimi. Numëron zgjedhjet dhe shfaq statistikat në fund.")

# 3. SI FUNKSIONON ZGJIDHJA
doc.add_heading("3. SI FUNKSIONON ZGJIDHJA", level=2)

doc.add_paragraph("1. Deklarimi i enum-it dhe variablave për numërues.")
doc.add_paragraph("2. Cikël for për tre iteracione.")
doc.add_paragraph("3. Leximi i input-it dhe validimi.")
doc.add_paragraph("4. Switch për reagimin sipas zgjedhjes.")
doc.add_paragraph("5. Përditësimi i numëruesve.")
doc.add_paragraph("6. Shfaqja e statistikave.")

# Save the document
doc.save("day3_task5_documentation.docx")